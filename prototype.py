#!/usr/bin/env python3.6

import sys, requests, html5lib, re
from bs4 import BeautifulSoup
from docx import Document

# Exit with error code 2 if correct arguments not provided
if len(sys.argv) != 3:
    print('ERROR: Need 2 arguments: 1, File name to write to. 2, Url to parse.')
    sys.exit(2)
else:
    filename = sys.argv[1]
    url = sys.argv[2]
    path = url.split('.com')[1]
    authority = url.split(path)[0]
    document = Document()

    # Parse website URL provided
    r = requests.get(url)
    soup = BeautifulSoup(r.text, 'html5lib')

    # Find all webpages for this site we want to content scrape
    internal_urls = []
    links = soup.find_all('a')
    for link in links:
        if link.has_attr('href'):
            href = link['href']
            if re.match(path, href):
                internal_urls.append(authority + href)

    unique_urls = []
    for url in internal_urls:
        if url not in unique_urls:
            unique_urls.append(url)

    # TODO add comment
    for url in unique_urls:
        # Parse website URL provided
        r = requests.get(url)
        soup = BeautifulSoup(r.text, 'html5lib')
        document.add_heading(url, 0)

        # Find page meta name and meta description
        # TODO Is there a more performant or succinct way
        metas = soup.find_all('meta')
        
        document.add_paragraph().add_run('Title:').bold = True
        for meta in metas:
            if 'name' in meta.attrs and meta.attrs['name'] == 'title':
                    document.add_paragraph(meta.attrs['content'])
                    break

        document.add_paragraph().add_run('Description:').bold = True
        for meta in metas:
            if 'name' in meta.attrs and meta.attrs['name'] == 'description':
                    document.add_paragraph(meta.attrs['content'])
                    break

        document.add_paragraph().add_run('Content:').bold = True

        # Stipulate all the body text we are interestd in 
        wanted_tags = ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li']

        # Find all divs that contain paragraph text
        p_divs = []
        p_tags = soup.find_all('p')
        for p in p_tags:
            p_divs.append(p.parent)

        # Add all info from the div to the document
        for div in p_divs:
            for tag in div.descendants:
                if tag.name in wanted_tags:
                    # If header tag add heading and type
                    if tag.name[0] == 'h':
                        header = f'{tag.text} - ({tag.name})'
                        document.add_heading(header, level=1)
                    # If ordered list item add number
                    # TODO will this increment number
                    elif tag.name == 'li' and tag.parent.name == 'ol':
                        document.add_paragraph(tag.text, style='List Number')
                    # If unordered list item add bullet point
                    elif tag.name == 'li' and tag.parent.name == 'ul':
                        document.add_paragraph(tag.text, style='List Bullet')
                    # Else paragraph, no formatting needed
                    else:
                        document.add_paragraph(tag.text)


        document.add_page_break()

    document.save(filename)
