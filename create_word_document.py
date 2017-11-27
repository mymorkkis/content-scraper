from docx import Document
from create_document import create_url_document


def create_word_document(filename, urls):
    word_doc = Document()
    # TODO add comment
    for url in urls:
        doc_section = create_url_document(url)
        word_doc.create_word_section(word_doc, doc_section)

    word_doc.save(filename)


def create_word_section(word_doc, doc_section):
    ds = doc_section
    word_doc.add_heading(ds.heading, 0)

    word_doc.add_paragraph().add_run('Title:').bold = True
    word_doc.add_paragraph(ds.title)

    word_doc.add_paragraph().add_run('Description:').bold = True
    word_doc.add_paragraph(ds.desciption)

    word_doc.add_paragraph().add_run('Content:').bold = True
    for item in ds.content:
        # Unpack tuple
        tag, content = item
        # If header tag add heading and type
        if tag[0] == 'h':
            header = f'{content} - ({tag})'
            word_doc.add_heading(header, level=1)
        # If unordered list item add bullet point
        elif tag == 'li':
            word_doc.add_paragraph(content, style='List Bullet')
        # Else paragraph, no formatting needed
        else:
            word_doc.add_paragraph(content)

    word_doc.add_page_break()
